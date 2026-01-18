import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =========================================================
# CONFIGURATION (EDIT PATHS ONLY IF NEEDED)
# =========================================================

# Base screenshots folder (contains timestamp subfolders)
BASE_SCREENSHOTS_DIR = r"F:\5Gcase\sre\maybank_wallet_withdrawal_prototype_express_ejs\screenshots"

# Input DOCX with placeholder <<AUTO_INSERT_SCREENSHOTS_HERE>>
INPUT_DOCX = r"F:\5Gcase\sre\maybank_wallet_withdrawal_prototype_express_ejs\doc\PartI_Scenario.docx"

# Output DOCX (generated)
OUTPUT_DOCX = r"F:\5Gcase\sre\maybank_wallet_withdrawal_prototype_express_ejs\doc\PartI_Scenario_WITH_Screenshots.docx"

# Screenshot width in inches (REQUIRED: 1.5")
IMAGE_WIDTH_INCHES = 1.5

# =========================================================
# HELPER FUNCTIONS
# =========================================================

def get_latest_screenshots_dir(base_dir):
    if not os.path.isdir(base_dir):
        raise FileNotFoundError(f"Screenshots base folder not found: {base_dir}")

    subdirs = [
        os.path.join(base_dir, d)
        for d in os.listdir(base_dir)
        if os.path.isdir(os.path.join(base_dir, d))
    ]

    if not subdirs:
        raise FileNotFoundError("No screenshot timestamp folders found.")

    return max(subdirs, key=os.path.getmtime)

def natural_sort_key(s):
    return [int(t) if t.isdigit() else t.lower() for t in re.split(r"(\d+)", s)]


def group_title_from_filename(filename: str) -> str:
    f = filename.lower()
    if "dashboard" in f:
        return "UC-008 Entry"
    if "reinvest" in f:
        # distinguish supporting UC-009 / UC-010 by keywords
        if "eligibility" in f or "not_eligible" in f or "not-eligible" in f:
            return "UC-009 Eligibility Validation Evidence"
        if "screening" in f or "blocked" in f:
            return "UC-010 Fraud & AML Screening Evidence"
        return "UC-008 Reinvest Funds Screens"
    if "misuse" in f:
        return "MUC-003 Fraudulent Reinvestment (Misuse Case)"
    if "security" in f:
        return "Security Mitigation Pages"
    if "admin_audit" in f or "admin" in f:
        return "UC-014 Audit Evidence"
    return "Other Screens"

def infer_page_name(filename: str) -> str:
    f = filename.lower()

    # Common UC-008 pages
    if "dashboard" in f:
        return "Dashboard (UC-008 Entry)"
    if "reinvest" in f and "confirm" in f:
        return "UC-008 – Reinvest (Confirm Details)"
    if "reinvest" in f and "success" in f:
        return "UC-008 – Reinvest Success (Receipt + Debit)"
    if "reinvest" in f and "insufficient" in f:
        return "UC-008 E1 – Insufficient Balance"
    if "eligibility" in f and ("noteligible" in f or "not_eligible" in f or "not-eligible" in f):
        return "UC-009 E2 – Not Eligible (Rejected)"
    if "not_eligible" in f or "not-eligible" in f:
        return "UC-009 E2 – Not Eligible (Rejected)"
    if "reinvest" in f and "eligibility" in f:
        return "UC-009 – Eligibility Check in Progress"
    if "reinvest" in f and "screening" in f:
        return "UC-010 – Fraud & AML Screening in Progress"
    if "reinvest" in f and "blocked" in f:
        return "UC-008 E3 – Blocked (Fraud/AML Failed) + Finance Alert"
    if "reinvest" in f:
        return "UC-008 – Reinvest Form (Product + Amount)"

    # Misuse
    if "misuse" in f:
        return "MUC-003 – Fraudulent Reinvestment Demonstration"

    # Security pages
    if "security_tips" in f or ("security" in f and "tips" in f):
        return "Security Mitigation – Security Tips"
    if "2fa" in f:
        return "Security Mitigation – Two-Factor Authentication (2FA)"
    if "linked" in f or "banks" in f:
        return "Security Mitigation – Manage Linked Banks"
    if "limits" in f:
        return "Security Mitigation – Manage Limits"
    if "report" in f:
        return "Security Mitigation – Report Suspicious Activity"

    # Audit
    if "admin_audit" in f or "audit" in f:
        return "UC-014 – Admin Audit Log (Evidence)"

    return "Prototype Screen"

def insert_paragraph_after(paragraph, text="", style=None):
    """
    XML-safe: insert a paragraph after the given one and return it.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para



def add_heading(doc, text):
    p = doc.add_paragraph(text)
    p.style = "Heading 3"

def add_centered_image(doc, image_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_path, width=Inches(IMAGE_WIDTH_INCHES))

def add_caption(doc, text, figure_no):
    p = doc.add_paragraph(f"Figure {figure_no}: {text}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True

# =========================================================
# MAIN EXECUTION
# =========================================================

def main():
    screenshots_dir = get_latest_screenshots_dir(BASE_SCREENSHOTS_DIR)

    images = sorted(
        [f for f in os.listdir(screenshots_dir) if f.lower().endswith(".png")],
        key=natural_sort_key
    )

    if not images:
        raise RuntimeError("No screenshots found in the screenshots folder.")

    print("Using screenshots folder:", screenshots_dir)
    print("Found screenshots:", len(images))

    doc = Document(INPUT_DOCX)

    # Locate placeholder
    insert_index = None
    for i, p in enumerate(doc.paragraphs):
        if "<<AUTO_INSERT_SCREENSHOTS_HERE>>" in p.text:
            insert_index = i
            p.text = ""
            break

    if insert_index is None:
        raise RuntimeError("Placeholder <<AUTO_INSERT_SCREENSHOTS_HERE>> not found in DOCX.")

    figure_no = 1

    for img in images:
        page_name = infer_page_name(img)
        img_path = os.path.join(screenshots_dir, img)

        # Insert page name
        heading_para = doc.add_paragraph(page_name)
        heading_para.style = "Heading 3"
        doc.paragraphs.insert(insert_index + 1, heading_para)

        # Insert image
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.add_run().add_picture(img_path, width=Inches(IMAGE_WIDTH_INCHES))
        doc.paragraphs.insert(insert_index + 2, img_para)

        # Insert caption
        caption_para = doc.add_paragraph(f"Figure {figure_no}: {page_name}")
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.runs[0].italic = True
        doc.paragraphs.insert(insert_index + 3, caption_para)

        # Spacer
        spacer = doc.add_paragraph("")
        doc.paragraphs.insert(insert_index + 4, spacer)

        figure_no += 1
        insert_index += 4

    doc.save(OUTPUT_DOCX)
    print(f"\n✅ DONE: Screenshots inserted and resized to 1.5\"")
    print(f"📄 Output file: {OUTPUT_DOCX}")

# =========================================================

if __name__ == "__main__":
    main()
